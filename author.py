"""
magiclib / author

------------------------------------------------------------------------------------------------------------------------
magiclib / author is a versatile Python library designed to simplify text and document processing. It provides tools for
analyzing, formatting, and managing text, Word documents, and PDF files, making it especially useful for academic and
professional work that requires consistency, automation, and efficiency. The library combines text analysis, document
formatting, and PDF manipulation into a single, user-friendly platform.
------------------------------------------------------------------------------------------------------------------------

Xinyuan Su
"""


import io
import os
import re
import sys
import time
import uuid
import PyPDF2
import shutil
import inspect
import difflib
import warnings
import platform
import requests
import subprocess
import contextlib
import pdfplumber
import numpy as np
import urllib.parse
from tqdm import tqdm
from PIL import Image
from docx import Document
from docx.shared import Pt, Cm
from sklearn.cluster import KMeans
from typing import Union, Tuple, Optional, List, Dict, Text
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT


Style = {
    "main_title": {
        "font": {
            "font_name": "Times New Roman",
            "font_size": 16,
            "bold": True
        },
        "paragraph": {
            "alignment": "center"  # 居中对齐
        }
    },
    "secondary_title": {
        "font": {
            "font_name": "Times New Roman",
            "font_size": 12,
            "bold": True
        },
        "paragraph": {
            "alignment": "justify"  # 两端对齐
        }
    },
    "tertiary_title": {
        "font": {
            "font_name": "Times New Roman",
            "font_size": 12,
            "bold": False
        },
        "paragraph": {
            "alignment": "justify"  # 两端对齐
        }
    },
    "body_text": {
        "font": {
            "font_name": "Times New Roman",
            "font_size": 12,
            "bold": False
        },
        "paragraph": {
            "alignment": "justify",  # 两端对齐
            "first_line_indent": 0.63  # 首行缩进 0.63 cm
        }
    },
    "figure_caption": {
        "font": {
            "font_name": "Times New Roman",
            "font_size": 12,
            "bold": False
        },
        "paragraph": {
            "alignment": "left",  # 左对齐
            "space_after": True  # 段后空一行
        }
    },
    "table_caption": {
        "font": {
            "font_name": "Times New Roman",
            "font_size": 12,
            "bold": False
        },
        "paragraph": {
            "alignment": "left",  # 左对齐
            "space_after": True  # 段后空一行
        }
    }
}


""" 检查或修改文本内容 """
class TextEditing:
    """
    主要用于英文文本内容。用于检查文本 (支持段落) 的内容，也可以用于修改指定内容
    Used to check the content of text (supporting paragraphs), and can also be used to modify specified content.
    Mainly used for English text content.

    注意： 接入的 data_dic 长度只允许为 1
    """

    # 初始化
    def __init__(self, text: Optional[Text] = None, text2: Optional[Text] = None):
        """
        # 接收参数 (2)
        :param text: (str) 需要处理的文本内容
        :param text2: (str) 需要处理的第二个文本内容
        """

        self.text = text
        self.text2 = text2

    # 检查文本内容的词数 / 字数
    def count_words(self, show_print: bool = True) -> Tuple[int, int]:
        """
        用于检查文本内容的词数 / 字数，中文则是为字数，英文则为词数
        The number of words/words used to check the content of the text is the number of words in Chinese
        and the number of words in English.

        注意：
        1.  英文会按照在一起的字母来检索，不会检查拼写是否正确
        2.  不会检查标点或者其它 emoji 符号
        Note:
        1.  English is searched by the letters together, without checking for correct spelling.
        2.  Don't check punctuation or other emojis.

        :param show_print: (bool) 是否打印结果，默认为 True

        :return num_english_words: (int) 英文本文内容的词数
        :return num_chinese_characters: (int) 中文本文内容的字数
        """

        # 检查赋值
        if self.text is not None:
            input_text = self.text
        else:
            class_name = self.__class__.__name__  # 获取类名
            method_name = inspect.currentframe().f_code.co_name  # 获取方法名
            raise ValueError(f"\033[95mIn {method_name} of {class_name}\033[0m, "
                             f"The variable text should be assigned a value and be a string.")

        # 统计中文字符数
        chinese_characters = re.findall(pattern='[\u4e00-\u9fff]', string=input_text)
        num_chinese_characters = len(chinese_characters)

        # 统计英文单词数
        english_words = re.findall(pattern='[a-zA-Z]+', string=input_text)
        num_english_words = len(english_words)

        # 打印结果
        if show_print:
            if num_chinese_characters > 0 and num_english_words > 0:
                print(
                    f"The input string contains \033[94m{num_english_words}\033[0m English words and "
                    f"\033[91m{num_chinese_characters}\033[0m Chinese characters.")
            elif num_chinese_characters > 0:
                print(f"The input string contains \033[91m{num_chinese_characters}\033[0m Chinese characters.")
            else:
                print(f"The input string contains \033[94m{num_english_words}\033[0m English words.")

        return num_english_words, num_chinese_characters

    # 比较文本内容
    def compare_text(self, show_print: bool = True) -> str:
        """
        比较两个文本内容，输出差异并高亮显示，并返回所有打印的内容。
        Compare the two text contents, print the difference and highlight it, and return all the printed content.

        注意：
        1.  根据中文和英文的句号来拆分句子，然后来判断每一句话是否一致
        2.  每个句号后会包容任意数量的与之相连的空白字符，即会忽略这些的存在
        Note:
        1. Split sentences according to the Chinese and English periods, and then determine whether
           each sentence is consistent.
        2. Any number of associated whitespace characters are ignored after each period.

        :param show_print: (bool) 是否打印结果，默认为 True

        :return output_text: (str) 包含所有需要打印内容的字符串
        """

        # 检查赋值
        if self.text is not None and self.text2 is not None:
            input_text1 = self.text
            input_text2 = self.text2
        else:
            class_name = self.__class__.__name__  # 获取类名
            method_name = inspect.currentframe().f_code.co_name  # 获取方法名
            raise ValueError(f"\033[95mIn {method_name} of {class_name}\033[0m, "
                             f"The variables text and text2 should be assigned values and both be a string.")

        # 使用正则表达式来根据中文和英文的句号拆分文本为句子
        sentences1 = re.split(pattern=r'(?<=[。.])\s*', string=input_text1)
        sentences1 = [s for s in sentences1 if s]

        sentences2 = re.split(pattern=r'(?<=[。.])\s*', string=input_text2)
        sentences2 = [s for s in sentences2 if s]

        # 使用difflib.Differ来比较两个句子列表
        d = difflib.Differ()
        diff = list(d.compare(sentences1, sentences2))

        result1 = []
        for line in diff:
            if line.startswith("- "):
                # 使用红色高亮显示第一个文本中独有的内容
                result1.append("\033[38;5;161m" + line[2:] + "\033[0m")
            elif line.startswith("  "):
                # 无差异的内容
                result1.append(line[2:])

        # 生成第二个文本的差异，这涉及到反转 + 和 - 符号
        reversed_diff = [line.replace('-', '+') if line.startswith('-')
                         else line.replace('+', '-') for line in diff]

        result2 = []
        for line in reversed_diff:
            if line.startswith("- "):
                # 使用蓝色高亮显示第二个文本中独有的内容
                result2.append("\033[34m" + line[2:] + "\033[0m")
            elif line.startswith("  "):
                # 无差异的内容
                result2.append(line[2:])

        # 打印第一个文本和差异
        output = ["-" * 50 + " Text 1 " + "-" * 50, "\n".join(result1), "-" * 50 + " Text 2 " + "-" * 50,
                  "\n".join(result2)]

        output_text = "\n".join(output)

        # 打印结果
        if show_print:
            print(output_text)

        # 返回所有打印的内容
        return output_text

    # 修正文本中空格的数量
    def normalize_spaces(self, show_print: bool = True) -> str:
        """
        根据中英文的要求来修正文本中空格的数量
        Normalize the number of Spaces in the text according to the requirements of Chinese and English.

        注意：
        1.  中文标点 '。！？，、；：' 后会直接接文本，而英文 '.!?,;:' 后会加一个空格再接文本
        2.  中文和英文之间会加上空格，无论哪个在前哪个在后
        3.  会删除开头和结尾的换行符，且删除多余一个的换行符使其保留一个
        Note:
        1.  Chinese punctuation '。！？，、；：' After will be directly followed by text,
            while English '.!?,;:' is followed by a space before the text.
        2.  A space is added between Chinese and English, whichever comes first and whichever comes last.
        3.  The beginning and ending newlines are deleted, and the extra newline character is deleted
            so that it remains one.

        :param show_print: (bool) 是否打印结果，默认为 True

        :return output_text: (str) 包含所有需要打印内容的字符串
        """

        # 检查赋值
        if self.text is not None:
            input_text = self.text
        else:
            class_name = self.__class__.__name__  # 获取类名
            method_name = inspect.currentframe().f_code.co_name  # 获取方法名
            raise ValueError(f"\033[95mIn {method_name} of {class_name}\033[0m, "
                             f"The variable text should be assigned a value and be a string.")

        # 移除文章开头的换行符
        text = re.sub(pattern=r'^\s+', repl='', string=input_text)

        # 连续应用正则表达式直到中文字符间的空格完全移除
        while re.search(pattern=r'([\u4e00-\u9fff])\s+([\u4e00-\u9fff])', string=text):
            text = re.sub(pattern=r'([\u4e00-\u9fff])\s+([\u4e00-\u9fff])', repl=r'\1\2',    string=text)

        # 在中文和英文之间添加必要的空格
        text = re.sub(pattern=r'([\u4e00-\u9fff])([A-Za-z])', repl=r'\1 \2', string=text)  # 中文后跟英文
        text = re.sub(pattern=r'([A-Za-z])([\u4e00-\u9fff])', repl=r'\1 \2', string=text)  # 英文后跟中文

        # 替换英文文本中多余的空格为单个空格
        text = re.sub(pattern=r'[ \t]+', repl=' ', string=text)

        # 处理所有中文标点，确保后面没有空格
        chinese_punctuation = "。！？，、；："
        for punct in chinese_punctuation:
            text = re.sub(pattern=r'\{}[\t ]*'.format(punct), repl=punct, string=text)

        # 处理所有英文标点，确保后面有一个空格
        english_punctuation = ".!?,;:"
        for punct in english_punctuation:
            text = re.sub(pattern=r'\{}[\t ]*'.format(punct), repl=punct + ' ', string=text)

        # 删除文本末尾不必要的空格及标点后的多余空格
        text = re.sub(pattern=r'\s+$', repl='', string=text)
        text = re.sub(pattern=r'\s+([.,;!?])', repl=r'\1', string=text)

        # 处理多余的换行符，只保留一个
        text_without_spaces = re.sub(pattern=r'\n+', repl='\n', string=text)

        # 打印结果
        if show_print:
            print(text_without_spaces)

        return text_without_spaces

    # 替换文本
    def replace_text(self, patterns_list: List[str], replacement: str = "", show_print: bool = True) -> str:
        """
        在 self.text 中将所有 patterns 列表里的子串替换成指定内容

        :param patterns_list: (list) 需要替换掉内容的 list
        :param replacement: (str) 需要替换成的内容
        :param show_print: (bool) 是否打印结果，默认为 True

        :return output_text: (str)  包含所有需要打印内容的字符串
        """

        # 检查赋值
        if self.text is not None:
            input_text = self.text
        else:
            class_name = self.__class__.__name__  # 获取类名
            method_name = inspect.currentframe().f_code.co_name  # 获取方法名
            raise ValueError(f"\033[95mIn {method_name} of {class_name}\033[0m, "
                             f"The variable text should be assigned a value and be a string.")

        for p in patterns_list:
            input_text = re.sub(p, replacement, input_text)

        output_text = input_text

        # 打印结果
        if show_print:
            print(output_text)

        return output_text


""" 对 word 文档进行修改 """
class Word:
    """
    用于对 word 文档的操作，查找，修改等
    It is used for operations, searching, modifying, etc. on word documents.

    格式需要符合外部变量 Style
    """

    Style = Style

    # 初始化
    def __init__(self, read_path: Optional[str] = None, save_path: Optional[str] = None):
        """
        接收参数

        :param read_path: (str) word 文档的读取路径
        :param save_path:  (str) word 文档的保存路径
        """

        self.read_path = read_path
        self.save_path = save_path

    # 对文本格式的修改
    def format_word_document(self, inspection_mode: bool = False) -> None:
        """
        修改 Word 文档中的内容格式，使用预定义的 self.Font_Style 字典
        Modify the format of content in a Word document, using the predefined self.Font_Style dictionary

        文章需要以下格式才能正常被识别：
        1.  开头第一段直接为文章的标题内容
        2.  图名必需以 'Fig.' 开头
        3.  表标题必需第一行以 'Table' 开头，且第二行为其名
        4.  正文在首个二级标题到参考文献之间

        The article must follow the format below to be properly recognized:
        1.	The first paragraph at the beginning should directly contain the title of the article.
        2.	Figure captions must start with “Fig.”
        3.	Table captions must begin with “Table” on the first line, followed by the name on the second line.
        4.  The main text should be located between the first second-level heading and the references section.

        :param inspection_mode: (bool) 是否打开检查模式，默认为 False

        :return: None
        """

        # 文件路径
        read_path = self.read_path
        save_path = self.save_path

        # 检查文件路径有效性
        if not os.path.isfile(self.read_path):
            class_name = self.__class__.__name__  # 获取类名
            method_name = inspect.currentframe().f_code.co_name  # 获取方法名
            raise ValueError(f"\033[95mIn {method_name} of {class_name}\033[0m, "
                             f"The specified read_path does not exist or is not a valid file: {self.read_path}")

        # 打开 Word 文档
        doc = Document(read_path)

        # 标志变量
        found_main_title = False
        in_body_text = False
        table_next_paragraph = False

        # 正则表达式模式
        second_level_heading_pattern = r"^\s*\d+\.\s*[^0-9]+$"
        third_level_heading_pattern = r"^\s*\d+\.\d+\.\s*[^0-9]+$"
        fig_title_pattern = r"^Fig\.\s*\d+\.\s*.*$"
        table_title_pattern = r"^Table\s*\d+\s*.*$"
        references_pattern = r"^References$"

        # 遍历文档中的每个段落
        for idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()  # 获取段落文本并去掉首尾空格
            if not text:  # 如果段落为空，则跳过
                continue

            # 打印原文段落
            if inspection_mode:
                print(f'{text}')

            # 如果还未找到主标题，则将当前段落处理为主标题
            if not found_main_title:
                run = para.runs[0] if para.runs else para.add_run()
                run.font.name = self.Style["main_title"]["font"]["font_name"]
                run.font.size = Pt(self.Style["main_title"]["font"]["font_size"])
                run.bold = self.Style["main_title"]["font"]["bold"]
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 使用正确的对齐方式
                if "first_line_indent" in self.Style["main_title"]["font"]:
                    para.paragraph_format.first_line_indent = Cm(self.Style["main_title"]["font"]["first_line_indent"])
                if ("space_after" in self.Style["main_title"]["font"]
                        and self.Style["main_title"]["font"]["space_after"]):
                    para.paragraph_format.space_after = Pt(12)
                found_main_title = True  # 标记已找到主标题
                print(f'\033[1m\033[3m{text}\033[0m')  # 打印主标题
                continue

            # 如果匹配到二级标题格式
            if re.match(second_level_heading_pattern, text):
                run = para.runs[0] if para.runs else para.add_run()
                run.font.name = self.Style["secondary_title"]["font"]["font_name"]
                run.font.size = Pt(self.Style["secondary_title"]["font"]["font_size"])
                run.bold = self.Style["secondary_title"]["font"]["bold"]
                para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 使用正确的对齐方式
                in_body_text = True  # 标记进入正文部分
                print(f'\033[91m{text}\033[0m')  # 打印二级标题
                continue

            # 如果匹配到三级标题格式
            if re.match(third_level_heading_pattern, text):
                run = para.runs[0] if para.runs else para.add_run()
                run.font.name = self.Style["tertiary_title"]["font"]["font_name"]
                run.font.size = Pt(self.Style["tertiary_title"]["font"]["font_size"])
                run.bold = self.Style["tertiary_title"]["font"]["bold"]
                para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 使用正确的对齐方式
                if inspection_mode:
                    print(f'\033[95m{text}\033[0m')  # 打印三级标题
                else:
                    print('', end='  ')
                    print(f'\033[95m{text}\033[0m')  # 打印三级标题
                continue

            # 如果匹配到图标题格式
            if re.match(fig_title_pattern, text):
                run = para.runs[0] if para.runs else para.add_run()
                run.font.name = self.Style["figure_caption"]["font"]["font_name"]
                run.font.size = Pt(self.Style["figure_caption"]["font"]["font_size"])
                run.bold = self.Style["figure_caption"]["font"]["bold"]
                para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 使用正确的对齐方式
                if inspection_mode:
                    print(f'\033[34;2m{text}\033[0m')  # 打印图标题
                continue

            # 如果匹配到表格标题格式
            if re.match(table_title_pattern, text):
                run = para.runs[0] if para.runs else para.add_run()
                run.font.name = self.Style["table_caption"]["font"]["font_name"]
                run.font.size = Pt(self.Style["table_caption"]["font"]["font_size"])
                run.bold = self.Style["table_caption"]["font"]["bold"]
                para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 使用正确的对齐方式
                table_next_paragraph = "(continued)" not in para.text  # 标记下一段也是表格标题
                if inspection_mode:
                    print(f'\033[32;2m{text}\033[0m')  # 打印表格标题
                continue

            # 如果上一段标记为表格标题的后续段落
            if table_next_paragraph:
                run = para.runs[0] if para.runs else para.add_run()
                run.font.name = self.Style["table_caption"]["font"]["font_name"]
                run.font.size = Pt(self.Style["table_caption"]["font"]["font_size"])
                run.bold = self.Style["table_caption"]["font"]["bold"]
                para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 使用正确的对齐方式
                table_next_paragraph = False  # 重置标记
                if inspection_mode:
                    print(f'\033[32;2m{text}\033[0m')  # 打印表格标题续
                continue

            # 如果匹配到参考文献部分
            if re.match(references_pattern, text):
                run = para.runs[0] if para.runs else para.add_run()
                run.font.name = self.Style["body_text"]["font"]["font_name"]
                run.font.size = Pt(self.Style["body_text"]["font"]["font_size"])
                run.bold = self.Style["body_text"]["font"]["bold"]
                para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 使用正确的对齐方式
                in_body_text = False  # 标记退出正文部分
                if inspection_mode:
                    print(f'\033[37;2m{text}\033[0m')  # 打印参考文献
                continue

            # 处理正文段落
            if (found_main_title and not re.match(second_level_heading_pattern, text) and in_body_text
                    and not re.match(fig_title_pattern, text) and not re.match(table_title_pattern, text)):
                run = para.runs[0] if para.runs else para.add_run()
                run.font.name = self.Style["body_text"]["font"]["font_name"]
                run.font.size = Pt(self.Style["body_text"]["font"]["font_size"])
                run.bold = self.Style["body_text"]["font"]["bold"]
                para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 使用正确的对齐方式
                if inspection_mode:
                    print(f'\033[4m{text}\033[0m')  # 打印正文段落）

        # 保存修改后的文档
        if save_path is not None:
            doc.save(save_path)

        return None


""" 对 PDF 进行操作 """
class PDF:
    """
    对 PDF 文件进行拼接、抽取、转换及压缩的功能
    The functions of splicing, extracting, converting and compressing PDF files.
    """

    # 初始化
    def __init__(self, read_path: Union[None, str, List[str]] = None, save_path: Optional[str] = None):
        """
        接收参数

        :param read_path: (str, list) 读取文件的路径，如果为 list 则为多个文件的路径，如果为 str 则为单一文件的路径
        :param save_path: (str) 生成文件的保存路径，为目录路径
        """

        self.read_path = read_path
        self.save_path = save_path

        # read_pdf()
        self.pdf_text_dict = {}

    # 将单个 PDF、Word 文档及图片合成一个长 PDF
    def merge_files_to_pdf(self, file_list: Optional[List] = None, save_path: str = None) -> None:
        """
        将多个文件（PDF、图片、Word 文档）合并为一个 PDF 文件。
        Merge multiple files (PDFs, images, and Word documents) into a single PDF.

        :param file_list: (list) 待合并的文件路径列表，支持 .pdf、图片（.jpg, .png 等）、Word 文档（.doc, .docx）
        :param save_path: (str) 合并后输出的 PDF 文件保存的目录路径

        支持的文件类型：
            - PDF 文件：直接加入合并器
            - 图片文件：转换为 PDF 后加入合并器（使用 PIL）
            - Word 文件：使用 LibreOffice 转换为 PDF 后加入合并器（需正确安装 LibreOffice）

        注意事项：
            - 需要安装 LibreOffice，下载: brew install --cask libreoffice
            - 用实际路径替换原本路径: which soffice
            - 临时生成的 PDF 文件在合并完成后会自动删除

        :return: 无返回值，合并结果直接保存为 save_path 目录下名为 merged_PDF.pdf 的文件
        """

        # 赋值处理
        if file_list is None:
            file_list = self.read_path
        if save_path is None:
            save_path = self.save_path

        # 检查赋值
        if file_list is None:
            class_name = self.__class__.__name__  # 获取类名
            method_name = inspect.currentframe().f_code.co_name  # 获取方法名
            raise ValueError(f"\033[95mIn {method_name} of {class_name}\033[0m, "
                            f"the parameter file_list cannot be empty. Please provide the list of files to be merged.")
        if save_path is None:
            class_name = self.__class__.__name__  # 获取类名
            method_name = inspect.currentframe().f_code.co_name  # 获取方法名
            raise ValueError(f"\033[95mIn {method_name} of {class_name}\033[0m, "
                             f"the parameter save_path cannot be empty. "
                             f"Please specify the path of the output directory.")

        # 检查 save_path 是否为有效目录
        class_name = self.__class__.__name__
        method_name = inspect.currentframe().f_code.co_name
        if not os.path.isdir(save_path):
            raise ValueError(f"\033[95mIn {method_name} of {class_name}\033[0m, "
                             f"The parameter save_path must be a valid directory: {save_path}")

        temp_pdfs = []  # 用于记录临时生成的 PDF 文件路径，以便后续删除
        merger = PyPDF2.PdfMerger()  # 创建 PDF 合并器对象
        print('')

        # 自定义进度显示，替换 tqdm
        for idx, file in enumerate(file_list, 1):
            file_name = os.path.basename(file).ljust(20)[:20]
            percent = int(idx / len(file_list) * 100)
            bar_length = 50
            filled_length = int(bar_length * percent // 100)
            bar = '\033[32;2m█\033[0m' * filled_length + '-' * (bar_length - filled_length)
            line = (f"\033[35;2m[{idx}/{len(file_list)}]\033[0m Being processed："
                    f"\033[36m{file_name}\033[0m | {bar} | \033[97m{percent}%\033[0m")
            sys.stdout.write('\r' + ' ' * (shutil.get_terminal_size().columns - 1) + '\r')
            sys.stdout.write(line)
            sys.stdout.flush()

            ext = os.path.splitext(file)[1].lower()  # 获取文件扩展名并转为小写

            # 如果是 PDF 文件，直接加入合并器
            if ext == '.pdf':
                try:
                    with contextlib.redirect_stderr(io.StringIO()):
                        merger.append(file)
                except Exception as e:
                    print(f"\nFailed to merge PDF: {file}, error: {e}")

            # 如果是图片文件，先转换为 PDF，再加入合并器
            elif ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff']:
                try:
                    image = Image.open(file).convert("RGB")
                    dir_name, base_name = os.path.split(file)
                    name, _ = os.path.splitext(base_name)
                    unique_suffix = str(uuid.uuid4())[:8]
                    pdf_path = os.path.join(dir_name, f"{name}_{unique_suffix}.pdf")
                    image.save(pdf_path)

                    if os.path.exists(pdf_path):
                        temp_pdfs.append(pdf_path)
                        with contextlib.redirect_stderr(io.StringIO()):
                            merger.append(pdf_path)
                    else:
                        print(f"\nNo PDF generated after image conversion: {file}")
                except Exception as e:
                    print(f"\nFailed to convert image: {file}, error: {e}")

            # 如果是 Word 文件，调用 LibreOffice 转换为 PDF
            elif ext in ['.doc', '.docx']:
                try:
                    dir_name, base_name = os.path.split(file)
                    name, _ = os.path.splitext(base_name)
                    pdf_path = os.path.join(dir_name, f"{name}.pdf")

                    result = subprocess.run([
                        "/opt/homebrew/bin/soffice",
                        "--headless",
                        "--convert-to", "pdf",
                        "--outdir", dir_name,
                        file
                    ], capture_output=True)

                    if os.path.exists(pdf_path):
                        temp_pdfs.append(pdf_path)
                        with contextlib.redirect_stderr(io.StringIO()):
                            merger.append(pdf_path)
                    else:
                        print(f"\nFailed to convert Word file to PDF: {file}")
                        print(result.stderr.decode())
                except Exception as e:
                    print(f"\nFailed to convert Word file: {file}, error: {e}")

            # 不支持的文件类型
            else:
                print(f"\nUnsupported file type: {file}")

        # 尝试将合并后的 PDF 写入指定输出路径（目录），文件名固定为 merged_PDF.pdf
        output_file = os.path.join(save_path, "merged_PDF.pdf")
        try:
            with open(output_file, 'wb') as f_out:
                merger.write(f_out)
            print(f"\nMerge completed, output file: \033[31m\033[1m{output_file}\033[0m")
        except OSError as e:
            print(f"\nFailed to write merged PDF: {e}")

        # 删除临时生成的 PDF 文件
        for temp_pdf in temp_pdfs:
            try:
                os.remove(temp_pdf)
            except OSError as e:
                print(f"\nFailed to delete temporary PDF: {temp_pdf}, error: {e}")

        return None

    # 从多个 PDF 文件中提取所有页面或指定页面并分别保存为新的 PDF 文件
    def extract_pages(self, file_list: Union[None, str, List[str]] = None, save_path: str = None,
                      page_numbers: Union[None, List[int], Tuple[int], range] = None) -> None:
        """
        从多个 PDF 文件中提取其所有页面，或指定页码，并将每个原始文件的页面分别保存为新的 PDF 文件
        Extracts all pages or specified pages from multiple PDF files and saves each as a new individual PDF.

        :param file_list: (list / str) 要处理的 PDF 文件路径列表。如果不传，则使用初始化时设置的 read_path
        :param save_path: (str) 提取后保存新 PDF 文件的目录路径。如果不传，则使用初始化时设置的 save_path
        :param page_numbers: (list / tuple) 要提取的页码，所有 PDF 均尝试提取相同的页码。如果某个 PDF 文件页数不足，将跳过处理并提示。
                             如果为 None，则为所有页

        注意事项：
            - 输出文件将保存在指定目录下，命名格式为 “原文件名_extracted_pages.pdf”
            - 如果指定路径无效、列表为空，或读取失败，会在终端输出相关错误信息

        :return: 无返回值，合并结果直接保存为 save_path 目录下名为 merged_PDF.pdf 的文件
        """

        # 如果未传入 file_list 或 save_path，则使用初始化时传入的值
        if file_list is None:
            file_list = self.read_path
        # 如果传入的是字符串路径，自动转为列表
        if isinstance(file_list, str):
            file_list = [file_list]
        if save_path is None:
            save_path = self.save_path

        # 如果 file_list 仍为 None，抛出异常
        if file_list is None:
            class_name = self.__class__.__name__  # 获取类名
            method_name = inspect.currentframe().f_code.co_name  # 获取方法名
            raise ValueError(f"\033[95mIn {method_name} of {class_name}\033[0m, "
                             f"the parameter file_list cannot be empty. Please provide the list of files to be merged.")
        # 如果 save_path 仍为 None，抛出异常
        if save_path is None:
            class_name = self.__class__.__name__  # 获取类名
            method_name = inspect.currentframe().f_code.co_name  # 获取方法名
            raise ValueError(f"\033[95mIn {method_name} of {class_name}\033[0m, "
                             f"the parameter save_path cannot be empty. "
                             f"Please specify the path of the output directory.")

        # 检查保存路径是否为有效目录
        class_name = self.__class__.__name__
        method_name = inspect.currentframe().f_code.co_name
        if not os.path.isdir(save_path):
            raise ValueError(f"\033[95mIn {method_name} of {class_name}\033[0m, "
                             f"The parameter save_path must be a valid directory: {save_path}")

        # 新增：收集成功和失败文件名
        success_files = []
        failed_files = []
        # 遍历每一个 PDF 文件，提取其全部页面或指定页面并另存为新 PDF
        for file in file_list:
            try:
                # 创建读取器对象，用于读取原 PDF 文件
                reader = PyPDF2.PdfReader(file)
                # 创建写入器对象，用于写入新 PDF 文件
                writer = PyPDF2.PdfWriter()

                if page_numbers is None:
                    # 提取全部页面
                    for page_num in range(len(reader.pages)):
                        writer.add_page(reader.pages[page_num])
                else:
                    # 检查最大页码是否超出范围
                    if max(page_numbers) > len(reader.pages):
                        failed_files.append(os.path.basename(file))
                        continue
                    for p in page_numbers:
                        # 页码从 0 开始
                        writer.add_page(reader.pages[p])

                # 构造输出文件名：原文件名_extracted_pages.pdf
                base_name = os.path.splitext(os.path.basename(file))[0]
                output_pdf = os.path.join(save_path, f"{base_name}_extracted_pages.pdf")

                # 写入新生成的 PDF 文件
                with open(output_pdf, 'wb') as output_file:
                    writer.write(output_file)

                # 成功，收集文件名
                success_files.append(os.path.basename(file))

            except (FileNotFoundError, PyPDF2.errors.PdfReadError, IndexError) as e:
                failed_files.append(os.path.basename(file))
                print(f"\033[90m[Error] {os.path.basename(file)}: {e}\033[0m")

        # 统一打印成功和失败的文件名
        if page_numbers is None:
            if success_files:
                print(f"\n\033[4mExtracted \033[32;2m\033[1mALL\033[0m\033[4m pages from:\033[0m")
                for i, name in enumerate(success_files, 1):
                    print(f"    {i}. \033[32m{name}\033[0m")
            if failed_files:
                print(f"\n\033[4mFailed to extract \033[31;2m\033[1mALL\033[0m\033[4m pages from:\033[0m")
                for i, name in enumerate(failed_files, 1):
                    print(f"    {i}. \033[31m{name}\033[0m")
        else:
            if success_files:
                print(f"\nExtracted pages \033[32;2m\033[1m{list(page_numbers)}\033[0m from:")
                for i, name in enumerate(success_files, 1):
                    print(f"    {i}. \033[32m{name}\033[0m")
            if failed_files:
                print(f"\nFailed to extract pages \033[31;2m\033[1m{list(page_numbers)}\033[0m from:")
                for i, name in enumerate(failed_files, 1):
                    print(f"    {i}. \033[31m{name}\033[0m")

        return None

    # 读取 PDF 的内容
    def read_pdf(self, pdf_path: Optional[str] = None, min_word_count: int = 0, clean_line: bool = False,
                 recursive: bool = False) -> dict:
        """
        从 PDF 文件或目录中提取正文文本，自动处理多栏排版
        Extract the main text from a PDF file or directory, automatically handling multi-column layouts.

        :param pdf_path: (str) PDF 文件路径或目录路径
        :param min_word_count: (int) 用于过滤短文本块，默认为 0
        :param clean_line: (bool) 是否清理换行符 '\n'，会去掉单个换行符，但也会使段落之间无换行，默认为 False
        :param recursive: (bool) 是否递归，即遍历目录下所有 PDF 文件，默认为 False

        :return pdf_text_dict: (dict) 读取的 PDF 文件的内容 {filename: extracted text}
        """

        # 忽略过时的代码警告
        warnings.simplefilter("ignore", FutureWarning)

        # 如果未指定路径，使用类属性
        if pdf_path is None:
            pdf_path = self.read_path

        pdf_files = []

        if os.path.isfile(pdf_path) and pdf_path.lower().endswith(".pdf"):
            pdf_files.append(pdf_path)

        elif os.path.isdir(pdf_path):
            if recursive:
                # 遍历目录及子目录
                for root, _, files in os.walk(pdf_path):
                    for file in files:
                        if file.lower().endswith(".pdf"):
                            pdf_files.append(os.path.join(root, file))
            else:
                # 只查当前目录
                for file in os.listdir(pdf_path):
                    full_path = os.path.join(pdf_path, file)
                    if os.path.isfile(full_path) and file.lower().endswith(".pdf"):
                        pdf_files.append(full_path)
        else:
            raise ValueError(f"It is not a valid PDF file or directory path: {pdf_path}")

        total_files = len(pdf_files)
        pdf_text_dict = {}
        processed_count = 0

        progress_bar = ""
        for idx, pdf_file in enumerate(pdf_files, 1):  # idx 从 1 开始
            page_texts = []
            try:
                with pdfplumber.open(pdf_file) as pdf:
                    for page in pdf.pages:
                        chars = page.chars
                        if not chars:
                            continue

                        # 自动检测列数（简单假设2栏，可根据需求改进）
                        x_coords = np.array([c["x0"] for c in chars]).reshape(-1, 1)
                        num_columns = 2  # 默认双栏，可改进为自动判断
                        if len(chars) < num_columns:
                            num_columns = 1
                        kmeans = KMeans(n_clusters=num_columns, random_state=0).fit(x_coords)
                        labels = kmeans.labels_

                        # 按列顺序拼接每页文本
                        columns_text = []
                        for col in sorted(np.unique(labels)):
                            col_chars = [c for i, c in enumerate(chars) if labels[i] == col]
                            # 列内按行排序
                            col_chars.sort(key=lambda c: c["top"])
                            # 合并每行
                            lines_dict = {}
                            for c in col_chars:
                                line_y = round(c["top"])
                                if line_y not in lines_dict:
                                    lines_dict[line_y] = []
                                lines_dict[line_y].append(c["text"])
                            col_text = "\n".join(["".join(line).strip() for y, line in sorted(lines_dict.items())
                                                  if len("".join(line).split()) >= min_word_count])
                            columns_text.append(col_text)

                        page_texts.append("\n".join(columns_text))

                # 拼合整份 PDF 的文本
                raw_text = "\n".join(page_texts)

                if clean_line:
                    # 连续三个及以上的 \n → 两个换行
                    clean_text = re.sub(pattern=r'\n{3,}', repl='\n\n', string=raw_text)
                    # 单个或非连续的 \n → 空格
                    clean_text = re.sub(pattern=r'(?<!\n)\n(?!\n)', repl=' ', string=clean_text)
                else:
                    clean_text = raw_text

                # 放入字典
                pdf_text_dict[os.path.basename(pdf_file)] = clean_text

            except Exception as e:
                print(f"\nError occurred when extracting {pdf_file}: {e}")
                pdf_text_dict[os.path.basename(pdf_file)] = ""

            # 更新进度条
            processed_count += 1
            progress_bar = "Processing PDFs \033[96m{}\033[0m / \033[94m{}\033[0m |{}{}| \033[35m{:.2f}%\033[0m".format(
                idx,
                total_files,
                "█" * int(processed_count * 50 / total_files),
                " " * (50 - int(processed_count * 50 / total_files)),
                processed_count / total_files * 100
            )
            print("\r" + progress_bar, end=" ")

        # 删除进度条
        print("\r" + " " * len(progress_bar) + "\r", end="")

        self.pdf_text_dict = pdf_text_dict
        return pdf_text_dict

    # 从默认打印机上打印文章
    def print_pdf(self, file_list: Union[None, str, List[str]] = None, show_result: bool = True) -> None:
        """
        从默认打印机上打印目标 PDF 文件
        Print the target PDF file from the default printer

        :param file_list: (list / str) 要打印的 PDF 文件路径列表。如果不传，则使用初始化时设置的 read_path
        :param show_result: (bool) 是否打印结果，无论是否成功都会受影响，默认为 True
        """

        # 如果未传入 file_list 或 save_path，则使用初始化时传入的值
        if file_list is None:
            file_list = self.read_path
        # 如果传入的是字符串路径，自动转为列表
        if isinstance(file_list, str):
            file_list = [file_list]

        system_name = platform.system()

        for pdf_path in file_list:
            # 检查文件是否存在
            if not os.path.exists(pdf_path):
                if show_result:
                    print(f"File does not exist, skipped: \033[31m{pdf_path}\033[0m")
                continue

            # 确认是否为 PDF 文件
            if not pdf_path.lower().endswith(".pdf"):
                if show_result:
                    print(f"Not a PDF file, skipped: \033[31m{pdf_path}\033[0m")
                continue

            try:
                if system_name == "Darwin":  # macOS
                    subprocess.run(
                        args=["lp", pdf_path],
                        check=True,
                        stdout=subprocess.DEVNULL  # 屏蔽标准输出
                    )

                elif system_name == "Windows":
                    # Windows 系统调用默认打印程序
                    os.startfile(pdf_path, "print")  # 必须是位置参数

                elif system_name == "Linux":
                    subprocess.run(
                        args=["lp", pdf_path],
                        check=True
                    )

                else:
                    if show_result:
                        print(f"Unsupported operating system: {system_name}")
                    continue

                if show_result:
                    print(f"Sent to printer: \033[34m{pdf_path}\033[0m")

            except Exception as e:
                print(f"Failed to print: {pdf_path}, Error: {e}")


""" 获取文章 / 资讯 """
class ArticleFetcher:
    """
    利用爬虫手段获取文章 / 成果 / 资讯
    Obtain articles/achievements/information by using web crawlers.
    """

    archaeology_journals = {
        # ==================== Elsevier ====================
        "Journal of Archaeological Science": {
            "eISSN": "1095-9238",
            "pISSN": "0305-4403"
        },
        "Journal of Archaeological Science: Reports": {
            "eISSN": "2352-409X",
            "pISSN": "2352-4103"
        },
        "Journal of Cultural Heritage": {
            "eISSN": "1778-3674",
            "pISSN": "1296-2074"
        },
        "Journal of Anthropological Archaeology": {
            "eISSN": "0278-4165",
            "pISSN": None
        },

        # ==================== Springer (BMC) ====================
        "Heritage Science": {
            "eISSN": "2050-7445",
            "pISSN": None
        },
        "International Journal of Historical Archaeology": {
            "eISSN": "1573-7748",
            "pISSN": "1092-7697"
        },

        # ==================== Taylor & Francis ====================
        "International Journal of Heritage Studies": {
            "eISSN": "1470-3576",
            "pISSN": None
        },
        "Australian Archaeology": {
            "eISSN": "2470-0363",
            "pISSN": "0312-2417"
        },
        "Journal of Field Archaeology": {
            "eISSN": "2042-4582",
            "pISSN": "0093-4690"
        },
        "World Archaeology": {
            "eISSN": "1470-1375",
            "pISSN": "0043-8243"
        },
        "Journal of Island and Coastal Archaeology": {
            "eISSN": "1556-4894",
            "pISSN": "1556-1828"
        },
        "Lithic Technology": {
            "eISSN": "2051-6185",
            "pISSN": "0197-7261"
        },
        "PaleoAmerica": {
            "eISSN": "2055-5571",
            "pISSN": "2055-5563"
        },
        "Journal of Conflict Archaeology": {
            "eISSN": "1574-0781",
            "pISSN": "1574-0773"
        },
        "Journal of Wetland Archaeology": {
            "eISSN": "1473-2976",
            "pISSN": None
        },
        "Palestine Exploration Quarterly": {
            "eISSN": "1743-1301",
            "pISSN": "0031-0328"
        },

        # ==================== Wiley ====================
        "Geoarchaeology": {
            "eISSN": "1520-6548",
            "pISSN": "0883-6353"
        },
        "Archaeometry": {
            "eISSN": "1475-4754",
            "pISSN": "0003-813X"
        },
        "Archaeology in Oceania": {
            "eISSN": "1834-4453",
            "pISSN": "0728-4896"
        },

        # ==================== Cambridge University Press ====================
        "American Antiquity": {
            "eISSN": None,
            "pISSN": "0002-7316"
        },
        "Antiquity": {
            "eISSN": "1745-1744",
            "pISSN": "0003-598X"
        },
        "Cambridge Archaeological Journal": {
            "eISSN": "1474-0540",
            "pISSN": "0959-7743"
        },
        "Archaeological Dialogues": {
            "eISSN": "1478-2294",
            "pISSN": "1380-2038"
        },

        # ==================== Archaeopress ====================
        "Journal of Greek Archaeology": {
            "eISSN": "2059-4682",
            "pISSN": "2059-4674"
        },
        "Journal of Hellenistic Pottery": {
            "eISSN": "2399-1852",
            "pISSN": "2399-1844"
        },
        "Proceedings of the Seminar for Arabian Studies": {
            "eISSN": None,
            "pISSN": "0308-8421"
        },

        # ==================== Peeters ====================
        "Ancient Near Eastern Studies": {
            "eISSN": None,
            "pISSN": "0066-6635"
        },
        "Iranica Antiqua": {
            "eISSN": "1783-1482",
            "pISSN": "0021-0870"
        },

        # ==================== 其他国际期刊 ====================
        "South African Archaeological Bulletin": {
            "eISSN": "0038-1969",
            "pISSN": None
        },
        "Antiguo Oriente": {
            "eISSN": None,
            "pISSN": "1667-9202"
        },
        "The Archaeological Journal": {  # Royal Archaeological Institute
            "eISSN": "0066-5983",
            "pISSN": None
        },
        "Il Capitale Culturale: Studies on the Value of Cultural Heritage": {  # Università di Macerata
            "eISSN": "2039-2362",
            "pISSN": None
        },

        # ==================== 中国考古学期刊（中文） ====================
        "Acta Archaeologica Sinica": {  # 考古学报
            "eISSN": "0453-2902",
            "pISSN": None
        },
        "考古": {
            "eISSN": "0453-2899",
            "pISSN": None
        },
        "文物": {
            "eISSN": "0511-4772",
            "pISSN": None
        },
        "考古与文物": {
            "eISSN": "1000-7830",
            "pISSN": None
        },
        "江汉考古": {
            "eISSN": "1001-0327",
            "pISSN": None
        },
        "华夏考古": {
            "eISSN": "1001-9928",
            "pISSN": None
        },
        "故宫博物院院刊": {
            "eISSN": "0452-7402",
            "pISSN": None
        },
        "四川文物": {
            "eISSN": "1003-6962",
            "pISSN": None
        },
        "北方文物": {
            "eISSN": "1001-0483",
            "pISSN": None
        },
        "文物保护与考古科学": {
            "eISSN": "1005-1538",
            "pISSN": None
        },
        "中国国家博物馆馆刊": {
            "eISSN": "2095-1639",
            "pISSN": None
        },
        "东南文化": {
            "eISSN": "1001-179X",
            "pISSN": None
        },
        "中原文物": {
            "eISSN": "1003-1731",
            "pISSN": None
        },
        "敦煌研究": {
            "eISSN": "1000-4106",
            "pISSN": None
        },
    }
    # 该端点支持的 filter 列表（根据 API 错误信息整理）
    filter_whitelist = {
        "from-approved-date", "until-approved-date", "has-assertion", "from-print-pub-date",
        "until-deposit-date", "from-accepted-date", "has-authenticated-orcid", "from-created-date",
        "relation.object", "issn", "ror-id", "lte-award-amount", "until-online-pub-date",
        "group-title", "full-text.application", "until-created-date", "license.version",
        "from-deposit-date", "has-abstract", "from-awarded-date", "has-event", "from-approved-date",
        "funder", "assertion-group", "from-online-pub-date", "from-issued-date", "directory",
        "content-domain", "license.url", "from-index-date", "full-text.version", "full-text.type",
        "until-posted-date", "has-orcid", "has-archive", "type", "has-ror-id", "is-update",
        "until-event-start-date", "update-type", "from-pub-date", "has-license", "funder-doi-asserted-by",
        "isbn", "has-full-text", "doi", "orcid", "has-content-domain", "prefix", "until-event-end-date",
        "has-funder", "award.funder", "clinical-trial-number", "member", "has-domain-restriction",
        "until-accepted-date", "container-title", "license.delay", "from-posted-date", "has-affiliation",
        "from-update-date", "has-award", "until-print-pub-date", "from-event-start-date",
        "gte-award-amount", "has-funder-doi", "has-alias", "until-index-date", "has-update",
        "until-update-date", "has-prime-doi", "until-issued-date", "until-pub-date", "award.number",
        "has-references", "type-name", "has-relation", "alternative-id", "archive", "relation.type",
        "updates", "relation.object-type", "category-name", "until-awarded-date",
        "has-clinical-trial-number", "assertion", "article-number", "has-update-policy",
        "from-event-end-date"
    }

    # 初始化
    def __init__(self, save_path: Optional[str] = None):
        """
        :param save_path: (str) 下载的保存路径
        """

        # 参数
        self.save_path = save_path

        # seek_doi()
        self.issn_list = []
        self.doi_list = []
        self.title_list = []

    # 获取目标期刊下的文章 DOI
    def seek_doi(self, issn_list: list = None, number: int = 500, query: Optional[str] = None,
                 show_result: bool = True, sort: str = "published", order: str = "desc",
                 **kwargs) -> Tuple[List[str], list[str]]:
        """
        获取目标期刊下的文章 DOI，支持过滤和模糊搜索。

        :param issn_list: 期刊 ISSN 列表（支持 eISSN 或 pISSN），若为 None 则默认使用 Journal of Archaeological Science 的
                          eISSN "1095-9238"
        :param number: 每个期刊返回的最大文章数（默认 500，最大 1000）
        :param query: 全局搜索关键词，会在标题、作者、摘要、机构等字段中模糊匹配
        :param show_result: 是否在控制台打印每篇文章的标题、DOI 和发表日期
        :param sort: 排序字段，可选 'published'（出版日期）、'issued'（发行日期）、'deposited'（入库日期）等，默认 'published'
        :param order: 排序顺序，'asc'（升序）或 'desc'（降序），默认 'desc'
        :param kwargs: 其他筛选参数，支持两类：
        1. filter 参数（精确过滤）：必须为白名单（FILTER_WHITELIST）中的字段，用于精确条件筛选。
           常用字段如下：
           - from_pub_date: 起始发表日期，格式 YYYY-MM-DD，例如 from_pub_date="2023-01-01"
           - until_pub_date: 截止发表日期，格式 YYYY-MM-DD，例如 until_pub_date="2023-12-31"
           - from_deposit_date: 起始入库日期（Crossref 收到元数据的日期）
           - until_deposit_date: 截止入库日期
           - from_online_pub_date: 起始在线出版日期
           - until_online_pub_date: 截止在线出版日期
           - from_issued_date: 起始发行日期（与 pub_date 类似，但更通用）
           - until_issued_date: 截止发行日期
           - type: 文献类型，常用值 journal-article, book-chapter, proceedings-article, dissertation 等
           - has_full_text: 是否有全文链接，值为 true 或 false
           - has_abstract: 是否有摘要，值为 true 或 false
           - has_orcid: 是否有 ORCID 标识，值为 true 或 false
           - has_affiliation: 是否包含作者机构信息，值为 true 或 false
           - has_funder: 是否有基金信息，值为 true 或 false
           - has_license: 是否有许可证信息，值为 true 或 false
           - license.url: 许可证 URL，支持通配符，例如 license.url:* 表示任何许可证
           - is_referenced_by_count: 被引次数，支持表达式如 >10, <5, 10-20
           - isbn: ISBN 号（用于图书）
           - issn: 期刊 ISSN（精确匹配，通常用于过滤）
           - member: 出版社成员 ID（数字），例如 member="78"（Elsevier）
           - container_title: 期刊名称（精确匹配）
           - category_name: 学科分类名称
           - archive: 存档机构
           - update_type: 更新类型
           - article_number: 文章编号
           - clinical_trial_number: 临床试验号
           - award.number: 基金项目编号
           - award.funder: 基金资助机构
           - orcid: ORCID 标识（精确匹配）
           - doi: DOI 号（精确匹配）
           - prefix: DOI 前缀（例如 "10.1016"）
           - ... 更多字段见 Crossref API 文档

        2. query 参数（模糊搜索）：自动添加 query. 前缀，用于在特定字段中模糊匹配。
           常用字段如下：
           - author: 作者姓名（支持部分匹配），自动转换为 query.author，例如 author="Smith"
           - title: 文章标题关键词，自动转换为 query.title，例如 title="bronze age"
           - affiliation: 作者机构名称，自动转换为 query.affiliation，例如 affiliation="Peking University"
           - container_title: 期刊名称（模糊匹配），自动转换为 query.container-title，例如 container_title="Archaeology"
           - abstract: 摘要关键词，自动转换为 query.abstract（注意：此参数仅对部分记录有效）
           - publisher: 出版社名称，自动转换为 query.publisher
           - funder: 基金资助机构，自动转换为 query.funder
           - alternative_id: 替代标识符（如出版商自己的 ID），自动转换为 query.alternative-id
           - 也可以显式使用 query_author、query_title 等，效果相同。

        3. 全局 query 参数：由独立参数 query 提供，用于在所有元数据字段中模糊搜索，优先级低于特定字段的 query 参数。
           - query: 自由文本，例如 query="Yangshao pottery"

        示例：
          # 按作者和日期过滤
          fetcher.seek_doi(issn_list=["1095-9238"], from_pub_date="2023-01-01", author="Wen Rui")
          # 按被引次数和全文过滤
          fetcher.seek_doi(has_full_text=True, is_referenced_by_count=">10")
          # 标题精确搜索+全局搜索
          fetcher.seek_doi(query_title="settlement", query="bronze age")

            注意：若传入的参数名不在白名单中，且不以 query- 开头，则自动视为 query 参数并添加 query. 前缀。

        :return: (tuple) 所有获取到的 DOI list 与标题 list

        示例：
            # 搜索 Journal of Archaeological Science 最近 100 篇有全文的文章
            fetcher.seek_doi(issn_list=["1095-9238"], number=100, has_full_text=True)

            # 搜索 2023 年后特定作者的文章
            fetcher.seek_doi(issn_list=["2352-409X"], from_pub_date="2023-01-01", author="Wen Rui")

            # 使用全局搜索和标题搜索
            fetcher.seek_doi(query="bronze age", query_title="settlement", number=50)

            # 按被引次数过滤
            fetcher.seek_doi(is_referenced_by_count=">10", type="journal-article")
        """

        if issn_list is None:
            issn_list = ["1095-9238"]
        elif not issn_list:
            issn_list = [ids['eISSN'] for ids in self.archaeology_journals.values() if ids.get('eISSN')]
        self.issn_list = issn_list

        start_time = time.time()
        all_dois = []
        all_titles = []

        for issn in issn_list:
            filter_parts = []
            query_parts = []

            # 处理 kwargs 中的参数
            for key, value in kwargs.items():
                if value is None:
                    continue
                crossref_key = key.replace("_", "-")
                if isinstance(value, bool):
                    value = str(value).lower()
                encoded_value = urllib.parse.quote(str(value))

                # 白名单内的参数作为 filter，否则作为 query 参数
                if crossref_key in self.filter_whitelist:
                    filter_parts.append(f"{crossref_key}:{encoded_value}")
                else:
                    # 若已经以 query- 开头，则直接转为 query. 格式
                    if crossref_key.startswith("query-"):
                        query_param = crossref_key.replace("-", ".")
                    else:
                        query_param = f"query.{crossref_key}"
                    query_parts.append(f"{query_param}={encoded_value}")

            # 处理显式的全局 query
            if query:
                encoded_query = urllib.parse.quote(str(query))
                query_parts.append(f"query={encoded_query}")

            filter_str = ",".join(filter_parts)
            query_str = "&".join(query_parts)

            # 查找期刊名称（用于打印）
            found = False
            for journal, ids in self.archaeology_journals.items():
                if issn == ids.get("eISSN") or issn == ids.get("pISSN"):
                    found = True
                    if show_result:
                        print(f"\033[33mArticles in {journal}:\033[0m")
            if not found and show_result:
                print(f"\033[31mJournal not found for ISSN {issn}\033[0m")

            # 构建 URL
            base_url = f"https://api.crossref.org/journals/{issn}/works?rows={number}&sort={sort}&order={order}"
            if filter_str:
                base_url += f"&filter={filter_str}"
            if query_str:
                base_url += f"&{query_str}"

            # 发送请求
            response = requests.get(base_url)
            if response.status_code == 200:
                data = response.json()
                articles = data['message']['items']

                for i, article in enumerate(articles):
                    doi = article.get('DOI', 'No DOI')
                    title = article.get('title', ['No Title'])[0]
                    all_dois.append(doi)
                    all_titles.append(title)

                    # 尝试获取出版日期（优先印刷版，其次在线版）
                    pub_date_parts = \
                        article.get('published-print',
                                    article.get('published-online', {})).get('date-parts', [[None]])[0]
                    pub_date = "-".join(str(p) for p in pub_date_parts if p is not None)

                    if show_result:
                        print(f"\033[36m{i + 1}.\033[0m \033[95m{title}\033[0m")
                        print(f"    DOI: https://doi.org/{doi}")
                        print(f"    Published: {pub_date}\n")
            else:
                if show_result:
                    print(f"\033[31mRequest failed\033[0m for ISSN \033[35m{issn}\033[0m:", response.status_code)
                    if response.status_code == 400:
                        print("Response content:", response.text)
                    print('')

        end_time = time.time()
        if show_result:
            print(f'Time-consuming: \033[36m{end_time - start_time}\033[0m')

        self.doi_list = all_dois
        self.title_list = all_titles
        return all_dois, all_titles
